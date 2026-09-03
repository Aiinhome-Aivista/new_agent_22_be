import json
from flask import Blueprint, request, jsonify, Response
from db import execute_query, execute_write

techlead_bp = Blueprint('techlead', __name__)

@techlead_bp.route('/validations', methods=['GET'])
def get_validations():
    track_id = request.args.get('track_id')
    
    query = """
        SELECT v.*, r.application_id, r.request_name, r.track_id 
        FROM validation_results v
        JOIN generation_requests r ON v.request_id = r.id
        WHERE (v.status = 'OPEN' OR v.status IS NULL)
    """
    params = []
    if track_id:
        query += " AND r.track_id = %s"
        params.append(track_id)
        
    query += """
        ORDER BY 
            CASE v.severity 
                WHEN 'error' THEN 1 
                WHEN 'warning' THEN 2 
                WHEN 'info' THEN 3 
                ELSE 4 
            END ASC,
            v.created_at DESC
    """
    results = execute_query(query, tuple(params))
    
    # Map severity for frontend
    severity_map = {
        'error': 'error',
        'warning': 'warning',
        'info': 'info'
    }
    
    for row in results:
        row['severity_display'] = severity_map.get(row.get('severity', 'info'), 'info')
        
    return jsonify({"success": True, "data": results})

@techlead_bp.route('/validations/<int:val_id>/action', methods=['POST'])
def action_validation(val_id):
    data = request.json or {}
    action = data.get('action') # 'WAIVE' or 'RESOLVE'
    
    if action not in ['WAIVE', 'RESOLVE']:
        return jsonify({"success": False, "message": "Invalid action"}), 400
        
    new_status = 'WAIVED' if action == 'WAIVE' else 'RESOLVED'
    
    execute_write("UPDATE validation_results SET status=%s WHERE id=%s", (new_status, val_id))
    
    return jsonify({"success": True, "message": f"Validation {new_status}"})

@techlead_bp.route('/validations/<int:val_id>/inspect', methods=['GET'])
def inspect_validation(val_id):
    # Fetch validation result
    val_res = execute_query("SELECT request_id, rule_name, message FROM validation_results WHERE id=%s", (val_id,))
    if not val_res:
        return jsonify({"success": False, "message": "Validation not found"}), 404
        
    req_id = val_res[0]['request_id']
    rule_name = val_res[0]['rule_name']
    message = val_res[0]['message']
    
    # Fetch generated files
    files = execute_query("SELECT file_name, file_content FROM generated_files WHERE request_id=%s", (req_id,))
    files_manifest = []
    if files:
        for f in files:
            files_manifest.append({
                "path": f['file_name'],
                "content": f['file_content']
            })
            
    import json
    from llm_service import call_llm, load_prompt
    
    prompt = load_prompt(
        "inspect_validation_prompt",
        rule_name=rule_name,
        message=message,
        files_manifest=json.dumps(files_manifest, default=str)
    )
    
    response_text = call_llm(prompt)
    
    try:
        import re
        # Clean up any potential markdown formatting
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
            
        response_text = re.sub(r'\\(?![/"\\bfnrtu])', r'\\\\', response_text)
        result = json.loads(response_text, strict=False)
        return jsonify({"success": True, "data": result})
    except Exception as e:
        print(f"Failed to parse LLM inspection response: {e}\nResponse: {response_text}")
        return jsonify({
            "success": True, 
            "data": {
                "excerpt": f"// Could not extract excerpt.\n// See file related to: {rule_name}",
                "suggestion": "Check the detailed validation report for more context."
            }
        })

@techlead_bp.route('/reviews', methods=['GET'])
def get_reviews():
    track_id = request.args.get('track_id')
    # Tech lead needs to review requests that have generated blueprints and code
    # typically status IN ('validated', 'packaged') waiting for final 'approved' (to git)
    # The architecture spec states TechLead reviews code for commit
    query = """
        SELECT r.id, r.request_name as serviceName, r.application_id as targetAppId, r.status, r.created_at as date, r.track_id
        FROM generation_requests r
        WHERE r.status IN ('validated', 'packaged', 'rework', 'approved', 'rejected')
    """
    params = []
    if track_id:
        query += " AND r.track_id = %s"
        params.append(track_id)
        
    query += " ORDER BY r.created_at DESC"
    results = execute_query(query, tuple(params))
    
    # Let's attach validation summary
    for req in results:
        req_id = req['id']
        val_rows = execute_query("SELECT * FROM validation_results WHERE request_id=%s AND passed=0", (req_id,))
        if not val_rows:
            req['validationStatus'] = '100% Passed'
        else:
            has_errors = any(v['severity'] == 'error' for v in val_rows)
            req['validationStatus'] = 'Failed' if has_errors else 'Passed with Warnings'
            
    return jsonify({"success": True, "data": results})

@techlead_bp.route('/reviews/signoff', methods=['POST'])
def signoff_review():
    data = request.json or {}
    request_id = data.get('request_id')
    decision = data.get('decision') # 'approved' or 'rework'
    comments = data.get('comments', '')
    reviewer_name = data.get('reviewer_name', 'Tech Lead')
    
    if decision not in ['approved', 'rework', 'rejected']:
        return jsonify({"success": False, "message": "Invalid decision"}), 400
        
    # Save review record
    execute_write(
        "INSERT INTO reviews (request_id, reviewer_name, decision, comments) VALUES (%s, %s, %s, %s)",
        (request_id, reviewer_name, decision, comments)
    )
    
    # Update request status
    # If TechLead approves code, it goes to 'approved' (ready for commit) or maybe 'completed'
    # According to our schema enum, we have 'approved', 'packaged', 'rework'
    execute_write("UPDATE generation_requests SET status=%s WHERE id=%s", (decision, request_id))
    
    if decision == 'approved':
        from agents.orchestrator_agent import start_packaging_thread
        start_packaging_thread(request_id)
    
    return jsonify({"success": True, "message": f"Sign-off recorded as {decision}"})

@techlead_bp.route('/reports/summary', methods=['GET'])
def get_reports_summary():
    track_id = request.args.get('track_id')
    
    # 1. total_passed: Count of requests with NO failed/warning validations
    # (Simplified: requests in packaged/approved state that have no error records)
    q1 = "SELECT id FROM generation_requests WHERE status IN ('validated', 'packaged', 'approved')"
    p1 = []
    if track_id:
        q1 += " AND track_id = %s"
        p1.append(track_id)
        
    total_reqs = execute_query(q1, tuple(p1))
    total_passed = 0
    for req in total_reqs:
        errors = execute_query("SELECT id FROM validation_results WHERE request_id=%s AND passed=0", (req['id'],))
        if not errors:
            total_passed += 1

    # 2. warnings_and_waivers: count of validation_results with severity in ('warning','info') or status='WAIVED'
    q2 = """SELECT COUNT(*) as count FROM validation_results v 
            JOIN generation_requests r ON v.request_id = r.id
            WHERE (v.severity IN ('warning','info') OR v.status='WAIVED')"""
    p2 = []
    if track_id:
        q2 += " AND r.track_id = %s"
        p2.append(track_id)
    res = execute_query(q2, tuple(p2))
    warnings = res[0]['count'] if res else 0

    # 3. critical_failures: count of OPEN error severity issues
    q3 = """SELECT COUNT(*) as count FROM validation_results v 
            JOIN generation_requests r ON v.request_id = r.id
            WHERE v.severity='error' AND (v.status='OPEN' OR v.status IS NULL)"""
    p3 = []
    if track_id:
        q3 += " AND r.track_id = %s"
        p3.append(track_id)
    res2 = execute_query(q3, tuple(p3))
    critical = res2[0]['count'] if res2 else 0

    return jsonify({
        "success": True, 
        "data": {
            "total_passed": total_passed,
            "warnings_and_waivers": warnings,
            "critical_failures": critical
        }
    })

@techlead_bp.route('/reports', methods=['GET'])
def get_reports_list():
    track_id = request.args.get('track_id')
    # Fetch completed or packaged generation requests as reports
    query = """
        SELECT id, request_name as title, application_id, 'PDF, DOCX' as type, created_at as date
        FROM generation_requests
        WHERE status IN ('validated', 'packaged', 'approved')
    """
    params = []
    if track_id:
        query += " AND track_id = %s"
        params.append(track_id)
        
    query += " ORDER BY created_at DESC"
    results = execute_query(query, tuple(params))
    
    for row in results:
        # Mocking size for UI
        row['size'] = f"{100 + (row['id'] * 15)} KB"
        # Format date for UI
        if row['date']:
            row['date'] = row['date'].strftime('%Y-%m-%d')
            
    return jsonify({"success": True, "data": results})

@techlead_bp.route('/reports/download/<int:report_id>', methods=['GET'])
def download_report(report_id):
    # Fetch full data to construct a JSON report
    req = execute_query("SELECT * FROM generation_requests WHERE id=%s", (report_id,))
    if not req:
        return jsonify({"success": False, "message": "Report not found"}), 404
        
    spec = execute_query("SELECT * FROM generation_specs WHERE request_id=%s", (report_id,))
    validations = execute_query("SELECT * FROM validation_results WHERE request_id=%s", (report_id,))
    
    report_data = {
        "report_metadata": {
            "id": report_id,
            "generated_on": str(req[0]['created_at']),
            "title": f"Audit Report for {req[0]['request_name']}"
        },
        "request_details": req[0],
        "specification": spec[0] if spec else {},
        "validations": validations or []
    }
    
    format_type = request.args.get('format', 'json').lower()
    
    if format_type == 'docx':
        from docx import Document
        import io
        
        doc = Document()
        doc.add_heading(report_data['report_metadata']['title'], 0)
        doc.add_paragraph(f"Generated On: {report_data['report_metadata']['generated_on']}")
        doc.add_heading('Request Details', level=1)
        for k, v in report_data['request_details'].items():
            doc.add_paragraph(f"{k}: {v}")
            
        doc.add_heading('Specification', level=1)
        for k, v in report_data['specification'].items():
            doc.add_paragraph(f"{k}: {v}")
            
        doc.add_heading('Validations', level=1)
        for v in report_data['validations']:
            passed_str = 'Yes' if v.get('passed') else 'No'
            doc.add_paragraph(f"[{str(v.get('severity', 'info')).upper()}] {str(v.get('rule_name', ''))}: {str(v.get('message', ''))} (Passed: {passed_str})")
            
        b = io.BytesIO()
        doc.save(b)
        b.seek(0)
        return Response(
            b.read(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-disposition": f"attachment; filename=audit_report_{report_id}.docx"}
        )
        
    elif format_type == 'pdf':
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        import io
        
        b = io.BytesIO()
        doc = SimpleDocTemplate(b, pagesize=letter)
        styles = getSampleStyleSheet()
        Story = []
        
        Story.append(Paragraph(report_data['report_metadata']['title'], styles['Title']))
        Story.append(Paragraph(f"Generated On: {report_data['report_metadata']['generated_on']}", styles['Normal']))
        Story.append(Spacer(1, 12))
        
        Story.append(Paragraph('Request Details', styles['Heading1']))
        for k, v in report_data['request_details'].items():
            Story.append(Paragraph(f"<b>{k}:</b> {str(v)}", styles['Normal']))
            
        Story.append(Spacer(1, 12))
        Story.append(Paragraph('Specification', styles['Heading1']))
        for k, v in report_data['specification'].items():
            Story.append(Paragraph(f"<b>{k}:</b> {str(v)}", styles['Normal']))
            
        Story.append(Spacer(1, 12))
        Story.append(Paragraph('Validations', styles['Heading1']))
        for v in report_data['validations']:
            passed_str = 'Yes' if v.get('passed') else 'No'
            Story.append(Paragraph(f"<b>[{str(v.get('severity', 'info')).upper()}]</b> {str(v.get('rule_name', ''))}: {str(v.get('message', ''))} (Passed: {passed_str})", styles['Normal']))
            Story.append(Spacer(1, 6))
            
        doc.build(Story)
        b.seek(0)
        return Response(
            b.read(),
            mimetype="application/pdf",
            headers={"Content-disposition": f"attachment; filename=audit_report_{report_id}.pdf"}
        )

    json_data = json.dumps(report_data, default=str, indent=4)
    
    return Response(
        json_data,
        mimetype="application/json",
        headers={"Content-disposition": f"attachment; filename=audit_report_{report_id}.json"}
    )
